import os
import ast
import csv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def get_module_docstring(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            tree = ast.parse(f.read())
            doc = ast.get_docstring(tree)
            return doc if doc else "Sem docstring no módulo."
    except Exception as e:
        return f"Erro ao ler docstring: {e}"

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)

def format_docstring_to_prose(docstring):
    # Simplifica o docstring para não ser apenas um raw text
    # Remove as linhas de "Autor" e formata os cabeçalhos.
    lines = docstring.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.startswith('='): continue
        if line.startswith('Autor:'): continue
        if line.startswith('Referência:'): continue
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines).strip()

def gerar_documento():
    doc = Document()
    
    # --- 1. CAPA ---
    doc.add_heading('Implementação e Análise de Estratégias de Algoritmos para o Problema da Mochila Binária (0/1 Knapsack)', 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\n[NOME COMPLETO - MATRÍCULA]').bold = True
    p.add_run('\n\nDisciplina: Análise e Projeto de Algoritmos\n\n\n')
    doc.add_page_break()
    
    # --- 2. INTRODUÇÃO ---
    doc.add_heading('2. Introdução', level=1)
    doc.add_paragraph(
        "O Problema da Mochila Binária (0/1 Knapsack Problem) é um clássico problema de otimização combinatória. "
        "Formalmente, é definido da seguinte maneira: dado um conjunto de n itens, onde cada item i possui um peso w_i e um valor v_i, "
        "e uma mochila com capacidade máxima W, o objetivo é encontrar um subconjunto de itens que maximize a soma dos valores "
        "sem que a soma dos pesos correspondentes exceda a capacidade W."
    )
    doc.add_paragraph(
        "A restrição binária (0/1) implica que cada item não pode ser fracionado: a decisão é estritamente binária, ou seja, "
        "o item é incluído integralmente na mochila ou é deixado de fora."
    )
    doc.add_heading('Aplicações Práticas', level=2)
    doc.add_paragraph(
        "1. Alocação de Orçamento e Investimentos: Em finanças corporativas, projetos representam os itens, "
        "o custo de implementação representa o peso, e o retorno financeiro esperado (ROI) representa o valor. A restrição 0/1 "
        "reflete o fato de que um projeto geralmente deve ser financiado integralmente para gerar retorno, e a capacidade da mochila "
        "é o capital total disponível da empresa para investimentos no trimestre."
    )
    doc.add_paragraph(
        "2. Logística de Carga e Transporte: No transporte aéreo ou marítimo, contêineres e caixas possuem pesos definidos "
        "e valores (que podem representar o custo do frete ou a prioridade da entrega). Um avião cargueiro possui um limite "
        "máximo de peso suportado. O problema se traduz em selecionar quais pacotes despachar para maximizar a receita do voo "
        "sem sobrecarregar a aeronave."
    )
    doc.add_paragraph(
        "3. Seleção de Processos em Sistemas Operacionais: Em escalonamento computacional sob restrições estritas de recursos, "
        "os itens são tarefas/processos, os pesos são a alocação necessária de RAM ou banda, e o valor é a prioridade ou urgência "
        "da tarefa para a estabilidade do sistema."
    )
    
    # --- 3. PROVA DE PERTINÊNCIA À CLASSE NP ---
    doc.add_heading('3. Prova de Pertinência à Classe NP', level=1)
    doc.add_paragraph(
        "Para discutir a complexidade estrutural, precisamos avaliar a versão de DECISÃO do Problema da Mochila: "
        "'Dado um conjunto de n itens (com pesos e valores), uma capacidade W e um limite mínimo de valor K, "
        "existe algum subconjunto S de itens tal que a soma dos pesos de S seja <= W e a soma dos valores de S seja >= K?'"
    )
    doc.add_paragraph(
        "Para provar que essa versão de decisão pertence à classe NP (Nondeterministic Polynomial time), precisamos mostrar "
        "que, se uma resposta for 'Sim', existe um certificado que pode ser verificado em tempo polinomial determinístico."
    )
    doc.add_paragraph(
        "Seja o certificado C o próprio subconjunto S de itens escolhidos. O algoritmo verificador opera da seguinte forma:\n"
        "1. Itera sobre os itens em C, somando seus pesos.\n"
        "2. Itera sobre os itens em C, somando seus valores.\n"
        "3. Verifica se a soma dos pesos <= W e a soma dos valores >= K."
    )
    doc.add_paragraph(
        "Como o certificado contém no máximo n itens, e a soma e as comparações aritméticas são operações de tempo constante, "
        "o algoritmo verificador roda em tempo O(n). Logo, como um candidato à solução pode ser verificado em tempo "
        "polinomial, o Problema da Mochila Binária pertence a NP."
    )
    doc.add_paragraph(
        "Embora pertença a NP, o problema é sabidamente NP-Difícil. A demonstração rigorosa disso envolve uma redução "
        "em tempo polinomial a partir do problema Subset Sum (Soma de Subconjuntos). De acordo com Cormen et al. [1] e "
        "Ziviani [2], como Subset Sum é NP-Completo, e ele pode ser mapeado como um caso especial do Problema da Mochila "
        "(onde o peso de um item é igual ao seu valor, e K = W), a Mochila Binária herda a complexidade NP-Difícil."
    )
    
    # --- 4. IMPLEMENTAÇÃO ---
    doc.add_heading('4. Implementação dos Paradigmas', level=1)
    doc.add_paragraph("O código completo das seis abordagens está disponível no Apêndice A. Abaixo, descrevemos o embasamento de cada uma e apresentamos os trechos focais das lógicas.")

    algoritmos = [
        ("Força Bruta", "src/forca_bruta.py"),
        ("Backtracking", "src/backtracking.py"),
        ("Divisão e Conquista", "src/divisao_conquista.py"),
        ("Programação Dinâmica", "src/prog_dinamica.py"),
        ("Algoritmo Guloso", "src/guloso.py"),
        ("Heurística FPTAS", "src/heuristica.py")
    ]

    for nome, arquivo in algoritmos:
        doc.add_heading(f"4.{algoritmos.index((nome, arquivo)) + 1} {nome}", level=2)
        docstring_puro = get_module_docstring(arquivo)
        prosa = format_docstring_to_prose(docstring_puro)
        doc.add_paragraph(prosa)
        
        # Extrações específicas:
        if nome == "Backtracking":
            doc.add_paragraph(
                "A implementação de Backtracking realiza duas podas cruciais para otimizar a árvore de decisão:\n"
                "1. Poda de Viabilidade: Interrompe a expansão de um nó caso o peso acumulado dos itens já selecionados ultrapasse a capacidade W.\n"
                "2. Poda de Limite (Bound): Utiliza uma relaxação gulosa fracionária para estimar o valor máximo alcançável a partir do nó atual. "
                "Se esse valor máximo projetado (bound) for menor ou igual à melhor solução já encontrada globalmente, o ramo é podado, evitando processamento inútil."
            )
        elif nome == "Divisão e Conquista":
            doc.add_paragraph(
                "A Divisão e Conquista estabelece um elo conceitual claro. Ela expressa a árvore de decisão da Força Bruta "
                "por meio do formalismo recursivo de subproblemas (incluir o item T(n-1, W-w_i) vs não incluir T(n-1, W)), "
                "mas, ao não aplicar memoização, acaba sofrendo da mesma repetição de estados e custo exponencial."
            )
        elif nome == "Algoritmo Guloso":
            doc.add_paragraph(
                "Instância Adversarial que prova a subotimalidade: Suponha uma capacidade W=50. Item A tem peso 51 e valor 100 (inviável). "
                "Item B tem peso 10 e valor 20 (razão 2.0). Item C tem peso 40 e valor 79 (razão 1.975). "
                "O Guloso escolhe o B pela razão, restando 40 de espaço. O C entra perfeitamente. Porém, se a capacidade fosse W=40, "
                "o Guloso escolheria B e ficaria com 30 de espaço (C não caberia mais), obtendo valor 20. A solução ótima "
                "seria escolher diretamente o C, obtendo valor 79. Essa falha ocorre pois o Guloso não prevê o 'buraco' que deixa na capacidade."
            )
        elif nome == "Heurística FPTAS":
            doc.add_paragraph(
                "Para garantir (1 - ε) do ótimo, o FPTAS sacrifica precisão nos bits menos significativos dos valores, "
                "aplicando um fator de escala K e arredondando para baixo. Isso reduz o teto do valor máximo da PD de O(n * V_max) "
                "para um polinômio limitado pelo inverso de ε. O resultado é um trade-off ajustável perfeito entre tempo e garantia matemática."
            )
            
    # --- 5. ANÁLISE DE COMPLEXIDADE ---
    doc.add_heading('5. Análise de Complexidade', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Algoritmo'
    hdr_cells[1].text = 'Complexidade de Tempo'
    hdr_cells[2].text = 'Complexidade de Espaço'
    
    complexidades = [
        ("Força Bruta", "O(2ⁿ)", "O(n)"),
        ("Backtracking", "O(2ⁿ) no pior caso", "O(n)"),
        ("Divisão e Conquista", "O(2ⁿ)", "O(n)"),
        ("Programação Dinâmica", "O(n × W)", "O(n × W) ou O(W)"),
        ("Guloso", "O(n log n)", "O(n)"),
        ("FPTAS", "O(n³ / ε)", "O(n² / ε)")
    ]
    for alg, tempo, espaco in complexidades:
        row_cells = table.add_row().cells
        row_cells[0].text = alg
        row_cells[1].text = tempo
        row_cells[2].text = espaco
        
    doc.add_paragraph(
        "\nJustificativas Matemáticas:\n"
        "- A Força Bruta e D&C avaliam explicitamente (ou implicitamente na recursão) duas escolhas para cada um dos n itens, totalizando 2^n folhas.\n"
        "- A Programação Dinâmica preenche uma matriz retangular onde as linhas são os n itens e as colunas as capacidades parciais de 0 até W. O preenchimento iterativo custa n*W operações, sendo pseudo-polinomial.\n"
        "- O Algoritmo Guloso é dominado pela operação de ordenação prévia das razões valor/peso, que utiliza Timsort no Python com custo limitante O(n log n).\n"
    )

    # --- 6. ANÁLISE DE RESULTADOS ---
    doc.add_heading('6. Análise de Resultados', level=1)
    
    doc.add_heading('6.1 Explosão Combinatória', level=2)
    doc.add_paragraph("Na análise de escalabilidade, a Força Bruta desponta absurdamente atingindo 2.784 segundos para N=16. Em oposição, a Programação Dinâmica resolve a mesma instância em 0.0007 segundos. O gráfico abaixo comprova o limiar da viabilidade prática dos exponenciais.")
    try:
        doc.add_picture('resultados/graficos/grafico_tempo_exponenciais.png', width=Inches(6.0))
        doc.add_picture('resultados/graficos/grafico_instrucoes_exponenciais.png', width=Inches(6.0))
    except Exception as e:
        doc.add_paragraph(f"[Erro ao carregar gráficos: {e}]")

    doc.add_heading('6.2 Eficácia da Poda no Backtracking', level=2)
    doc.add_paragraph("O Backtracking se provou capaz de suprimir grande parte do espaço de busca combinatório. Para N=16, o espaço total era de 131.071 nós, e a árvore com podas processou apenas 344 nós (uma varredura de apenas 0.26%). Isso se converteu em uma redução temporal que deixou a execução quase instantânea frente à Força Bruta pura.")
    try:
        doc.add_picture('resultados/graficos/grafico_eficacia_poda.png', width=Inches(6.0))
    except Exception: pass

    doc.add_heading('6.3 Escalabilidade dos Polinomiais', level=2)
    doc.add_paragraph("Testando para capacidades maiores (N=100 e W superior a 16.000), a Programação Dinâmica manteve uma progressão estruturada pseudo-polinomial (cerca de 0.18s a 1.0s). O FPTAS demorou mais (cerca de 6.8s) em virtude do fator matemático que reverte o ganho do espaço quando ε é pequeno (0.1).")
    try:
        doc.add_picture('resultados/graficos/grafico_tempo_escalaveis.png', width=Inches(6.0))
    except Exception: pass
    
    doc.add_heading('6.4 Qualidade das Soluções (Trade-off)', level=2)
    doc.add_paragraph(
        "A avaliação das instâncias adversariais apontou que o Guloso apresentou GAP de 21.25% para a instância N=5 "
        "e 1.90% para N=10, validando seu comportamento subótimo. "
        "Já o FPTAS, utilizando epsilon flexível na instância N=25, manteve precisão absoluta (0.0% gap) "
        "para epsilons rígidos (0.01 consumindo 0.92s) e atingiu um GAP máximo tolerado de 20.57% quando epsilon foi afrouxado para 0.5 (consumindo ínfimos 0.02s)."
    )
    try:
        doc.add_picture('resultados/graficos/grafico_gap_guloso.png', width=Inches(6.0))
        doc.add_picture('resultados/graficos/grafico_gap_fptas_epsilon.png', width=Inches(6.0))
    except Exception: pass

    # --- 7. CONCLUSÃO ---
    doc.add_heading('7. Conclusão', level=1)
    doc.add_paragraph(
        "O estudo empírico confirmou as expectativas teóricas da complexidade computacional para o problema NP-Difícil da Mochila Binária. "
        "As estratégias de exploração exaustiva (Força Bruta e Divisão e Conquista sem memoização) são inviáveis na prática devido à "
        "explosão exponencial 2^n. "
    )
    doc.add_paragraph(
        "O Backtracking provou que o uso de inteligência analítica na forma de limites superiores (bound fractionário) é eficiente "
        "na filtragem da árvore de decisão. O destaque prático absoluto é a Programação Dinâmica, que demonstrou a melhor escalabilidade "
        "na faixa pseudo-polinomial, embora seu consumo de memória O(nW) demande cautela. Finalmente, heurísticas como Guloso e FPTAS "
        "demonstraram como ceder à subotimalidade em favor do determinismo computacional é uma escolha sólida de engenharia para inputs colossais."
    )
    
    # --- 8. BIBLIOGRAFIA ---
    doc.add_heading('8. Bibliografia', level=1)
    doc.add_paragraph("[1] CORMEN, Thomas H. et al. Introduction to Algorithms. 3. ed. MIT Press, 2009.")
    doc.add_paragraph("[2] ZIVIANI, Nivio. Projeto de Algoritmos com implementações em Pascal e C. 3. ed. Thomson, 2004.")
    doc.add_paragraph("[3] VAZIRANI, Vijay V. Approximation Algorithms. Springer, 2001.")

    # --- 9. APÊNDICE A ---
    doc.add_page_break()
    doc.add_heading('9. Apêndice A — Código-fonte Completo', level=1)
    
    arquivos_apendice = [
        "src/instrumentacao.py",
        "src/gerador_instancias.py",
        "src/forca_bruta.py",
        "src/divisao_conquista.py",
        "src/backtracking.py",
        "src/prog_dinamica.py",
        "src/guloso.py",
        "src/heuristica.py",
        "src/main.py",
        "src/analise.py"
    ]
    
    for arq in arquivos_apendice:
        doc.add_heading(arq.replace('src/', ''), level=2)
        try:
            with open(arq, 'r', encoding='utf-8') as f:
                linhas = f.readlines()
                codigo_numerado = "".join(f"{idx+1:03d} | {linha}" for idx, linha in enumerate(linhas))
                add_code_block(doc, codigo_numerado)
        except Exception as e:
            doc.add_paragraph(f"Erro ao carregar arquivo: {e}")

    doc.save('Documentacao_Mochila_Binaria.docx')
    print("Documento DOCX gerado com sucesso!")

if __name__ == "__main__":
    gerar_documento()
